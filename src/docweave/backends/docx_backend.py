"""Microsoft Word (.docx) backend using python-docx."""

from __future__ import annotations

import hashlib
import os
from pathlib import Path
from typing import Any

from docweave.backends.base import BackendAdapter
from docweave.models import Block, InspectResult, NormalizedDocument, SourceSpan


def _hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def _get_style_id(element: Any) -> str:
    """Read the style ID directly from the XML element, avoiding part lookups."""
    from docx.oxml.ns import qn

    p_pr = element.find(qn("w:pPr"))
    if p_pr is not None:
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            return p_style.get(qn("w:val"), "")
    return ""


def _has_num_pr(element: Any) -> bool:
    """Check if a paragraph element has numbering properties (list item)."""
    from docx.oxml.ns import qn

    p_pr = element.find(qn("w:pPr"))
    if p_pr is not None:
        return p_pr.find(qn("w:numPr")) is not None
    return False


def _get_paragraph_text(element: Any) -> str:
    """Extract plain text from a w:p element by concatenating all w:t elements."""
    from docx.oxml.ns import qn

    texts = []
    for t_elem in element.iter(qn("w:t")):
        if t_elem.text:
            texts.append(t_elem.text)
    return "".join(texts)


def _get_table_text(element: Any) -> str:
    """Extract cell text from a w:tbl element."""
    from docx.oxml.ns import qn

    rows = element.findall(qn("w:tr"))
    row_texts = []
    for row in rows:
        cells = row.findall(qn("w:tc"))
        cell_strs = []
        for cell in cells:
            # Get text from all paragraphs in cell
            parts = []
            for p in cell.findall(qn("w:p")):
                parts.append(_get_paragraph_text(p))
            cell_strs.append(" ".join(parts).strip())
        row_texts.append(" | ".join(cell_strs))
    return "\n".join(row_texts)


class WordBackend(BackendAdapter):
    @property
    def name(self) -> str:
        return "word-docx"

    @property
    def tier(self) -> int:
        return 1

    @property
    def extensions(self) -> set[str]:
        return {".docx"}

    def detect(self, path: Path) -> float:
        return 1.0 if path.suffix.lower() in self.extensions else 0.0

    def load_view(self, path: Path) -> NormalizedDocument:
        from docx import Document

        document = Document(str(path))
        body = document.element.body

        blocks: list[Block] = []
        heading_stack: list[str] = []
        seq = 0
        para_index = 0

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para_index += 1
                text = _get_paragraph_text(child).strip()
                style_id = _get_style_id(child)

                # Determine block kind from style ID
                if style_id.startswith("Heading"):
                    kind = "heading"
                    try:
                        level = int(style_id.replace("Heading", "").strip())
                    except (ValueError, IndexError):
                        level = 1

                    while len(heading_stack) >= level:
                        heading_stack.pop()
                    heading_stack.append(text)

                elif "List" in style_id or _has_num_pr(child):
                    kind = "list_item"
                    level = None
                else:
                    kind = "paragraph"
                    level = None

                # Skip empty non-heading paragraphs
                if not text and kind != "heading":
                    continue

                seq += 1
                blocks.append(Block(
                    block_id=f"blk_{seq:03d}",
                    kind=kind,
                    section_path=list(heading_stack),
                    text=text,
                    raw_text=text,
                    level=level if kind == "heading" else None,
                    source_span=SourceSpan(
                        start_line=para_index,
                        end_line=para_index,
                    ),
                    stable_hash=_hash(text),
                ))

            elif tag == "tbl":
                para_index += 1
                text = _get_table_text(child)

                seq += 1
                blocks.append(Block(
                    block_id=f"blk_{seq:03d}",
                    kind="table",
                    section_path=list(heading_stack),
                    text=text,
                    raw_text=text,
                    level=None,
                    source_span=SourceSpan(
                        start_line=para_index,
                        end_line=para_index,
                    ),
                    stable_hash=_hash(text),
                ))

        return NormalizedDocument(
            file=str(path),
            backend=self.name,
            blocks=blocks,
            metadata={"_docx_document": document},
        )

    def inspect(self, path: Path) -> InspectResult:
        doc = self.load_view(path)
        return InspectResult(
            file=str(path),
            backend=self.name,
            tier="native-safe",
            editable=os.access(path, os.W_OK),
            supports={
                "comments": True,
                "tables": True,
                "styles": True,
                "track_changes": False,
            },
            fidelity={
                "write_mode": "native",
                "roundtrip_risk": "medium",
            },
            block_count=doc.block_count,
            headings=[b.text for b in doc.blocks if b.kind == "heading"],
        )

    def resolve_anchor(self, view: Any, anchor: dict[str, Any]) -> Any:
        from docweave.anchors import Anchor
        from docweave.anchors import resolve_anchor as _resolve

        parsed = Anchor(**anchor) if isinstance(anchor, dict) else anchor
        return _resolve(view, parsed)

    def extract_text(self, path: Path) -> str:
        """Extract plain text by concatenating all paragraph texts."""
        from docx import Document

        document = Document(str(path))
        return "\n".join(p.text for p in document.paragraphs)

    def plan(self, view: Any, patches: list[dict[str, Any]]) -> dict[str, Any]:
        raise NotImplementedError("Use generate_plan() directly")

    def apply(self, view: Any, plan: dict[str, Any]) -> str:
        raise NotImplementedError("Use apply_plan_docx() directly")

    def validate(self, original: str, modified: str) -> list[dict[str, Any]]:
        raise NotImplementedError(
            "For Word files, use validate_document() on the parsed NormalizedDocument"
        )

    def diff(self, original: str, modified: str) -> dict[str, Any]:
        raise NotImplementedError(
            "For Word files, use semantic_diff() on parsed NormalizedDocuments"
        )
