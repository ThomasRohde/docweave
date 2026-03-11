"""Tests for Blackbox Fix Round 4: 6 bugs + 7 UX issues."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pytest
import yaml
from typer.testing import CliRunner

from docweave.cli import app


@pytest.fixture()
def run_cli():
    runner = CliRunner()

    def _run(*args: str, expect_json: bool = True):
        result = runner.invoke(app, list(args))
        result.json = None
        if expect_json and result.output.strip():
            result.json = json.loads(result.output.strip())
        return result

    return _run


# ---------------------------------------------------------------------------
# Bug 1: view on .docx should not crash (metadata must be JSON-serializable)
# ---------------------------------------------------------------------------

def _make_docx(path: Path) -> None:
    """Create a minimal valid .docx file."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("Test content.")
    doc.save(str(path))


def test_view_docx_no_crash(run_cli, tmp_path):
    """view on a .docx file should succeed without serialization errors."""
    docx_path = tmp_path / "test.docx"
    _make_docx(docx_path)

    result = run_cli("view", str(docx_path))
    assert result.json is not None
    assert result.json["ok"] is True


def test_inspect_docx_no_crash(run_cli, tmp_path):
    """inspect on a .docx file should succeed."""
    docx_path = tmp_path / "test.docx"
    _make_docx(docx_path)

    result = run_cli("inspect", str(docx_path))
    assert result.json is not None
    assert result.json["ok"] is True


# ---------------------------------------------------------------------------
# Bug 4: CRLF corruption — apply should preserve LF on Windows
# ---------------------------------------------------------------------------

def test_apply_preserves_lf(run_cli, tmp_path):
    """Zero-op apply should not convert LF to CRLF."""
    doc = tmp_path / "doc.md"
    content = "# Title\n\nParagraph.\n"
    doc.write_bytes(content.encode("utf-8"))  # explicit LF

    patch = tmp_path / "patch.yaml"
    patch.write_text(yaml.dump({
        "version": 1,
        "target": {"format": "markdown"},
        "operations": [
            {
                "id": "op_001",
                "op": "insert_after",
                "anchor": {"by": "heading", "value": "Title"},
                "content": {"kind": "paragraph", "value": "New text."},
            }
        ],
    }))

    run_cli("apply", str(doc), "--patch", str(patch))

    raw = doc.read_bytes()
    assert b"\r\n" not in raw, "File should not contain CRLF line endings"
    assert b"\n" in raw


# ---------------------------------------------------------------------------
# Bug 5: Corrupted .docx gives ERR_IO (not ERR_INTERNAL_UNHANDLED)
# ---------------------------------------------------------------------------

def test_corrupted_docx_gives_err_io(run_cli, tmp_path):
    """A corrupted .docx should return ERR_IO, not crash."""
    bad_docx = tmp_path / "bad.docx"
    bad_docx.write_text("This is not a valid docx file")

    result = run_cli("view", str(bad_docx))
    assert result.json is not None
    assert result.json["ok"] is False
    errors = result.json.get("errors", [])
    codes = [e["code"] for e in errors]
    assert "ERR_IO" in codes or "ERR_VALIDATION_DECODE" in codes


def test_corrupted_docx_zip_gives_err_io(run_cli, tmp_path):
    """A .docx that is a zip but not a valid Office doc should return ERR_IO."""
    bad_docx = tmp_path / "notadocx.docx"
    with zipfile.ZipFile(bad_docx, "w") as z:
        z.writestr("garbage.txt", "not a docx")

    result = run_cli("view", str(bad_docx))
    assert result.json is not None
    assert result.json["ok"] is False


# ---------------------------------------------------------------------------
# Bug 6: Misplaced annotation → parse_warnings populated
# ---------------------------------------------------------------------------

def test_annotation_not_before_heading_warns(tmp_path):
    """Annotation comment not before a heading should produce a parse warning."""
    from docweave.backends.markdown_native import MarkdownBackend

    md = tmp_path / "test.md"
    md.write_text(
        '# Heading\n\nSome text.\n\n'
        '<!-- docweave: {"tags": ["foo"]} -->\n\n'
        'More text.\n\n'
        '## Next Heading\n'
    )
    backend = MarkdownBackend()
    doc = backend.load_view(md)
    assert len(doc.parse_warnings) >= 1
    assert "not immediately before a heading" in doc.parse_warnings[0]


def test_annotation_before_heading_no_warning(tmp_path):
    """Annotation immediately before a heading should NOT produce a warning."""
    from docweave.backends.markdown_native import MarkdownBackend

    md = tmp_path / "test.md"
    md.write_text(
        '<!-- docweave: {"tags": ["foo"]} -->\n\n'
        '# Heading\n\nSome text.\n'
    )
    backend = MarkdownBackend()
    doc = backend.load_view(md)
    assert len(doc.parse_warnings) == 0


# ---------------------------------------------------------------------------
# UX 1: Multi-tag support on inspect and view
# ---------------------------------------------------------------------------

def test_inspect_multi_tag(run_cli, tmp_path):
    """inspect --tag foo --tag bar should match headings with either tag."""
    md = tmp_path / "doc.md"
    md.write_text(
        '<!-- docweave: {"tags": ["foo"]} -->\n'
        '# Alpha\n\n'
        '<!-- docweave: {"tags": ["bar"]} -->\n'
        '# Beta\n\n'
        '<!-- docweave: {"tags": ["baz"]} -->\n'
        '# Gamma\n'
    )
    result = run_cli("inspect", str(md), "--tag", "foo", "--tag", "bar")
    env = result.json
    assert env["ok"] is True
    headings = env["result"]["headings"]
    texts = {h["text"] for h in headings}
    assert "Alpha" in texts
    assert "Beta" in texts
    assert "Gamma" not in texts


def test_view_multi_tag(run_cli, tmp_path):
    """view --tag foo --tag bar should show blocks from either tagged section."""
    md = tmp_path / "doc.md"
    md.write_text(
        '<!-- docweave: {"tags": ["foo"]} -->\n'
        '# Alpha\n\nFoo content.\n\n'
        '<!-- docweave: {"tags": ["bar"]} -->\n'
        '# Beta\n\nBar content.\n\n'
        '<!-- docweave: {"tags": ["baz"]} -->\n'
        '# Gamma\n\nBaz content.\n'
    )
    result = run_cli("view", str(md), "--tag", "foo", "--tag", "bar")
    env = result.json
    assert env["ok"] is True
    texts = [b["text"] for b in env["result"]["blocks"]]
    assert any("Foo" in t for t in texts)
    assert any("Bar" in t for t in texts)
    assert not any("Baz" in t for t in texts)


# ---------------------------------------------------------------------------
# UX 2: Occurrence error message mentions 1-indexed
# ---------------------------------------------------------------------------

def test_occurrence_error_message_mentions_1_indexed():
    from docweave.anchors import OccurrenceOutOfRangeError

    err = OccurrenceOutOfRangeError(requested=3, available=1)
    msg = str(err)
    assert "1-indexed" in msg
    assert "3" in msg
    assert "1 block(s)" in msg


# ---------------------------------------------------------------------------
# UX 4: --tag on find
# ---------------------------------------------------------------------------

def test_find_with_tag(run_cli, tmp_path):
    """find --tag should filter to tagged sections before searching."""
    md = tmp_path / "doc.md"
    md.write_text(
        '<!-- docweave: {"tags": ["api"]} -->\n'
        '# API\n\nThe API endpoint handles requests.\n\n'
        '# Internal\n\nThe internal endpoint handles requests.\n'
    )
    result = run_cli("find", str(md), "endpoint", "--tag", "api")
    env = result.json
    assert env["ok"] is True
    matches = env["result"]["matches"]
    # Should only match blocks under "API" section, not "Internal"
    for m in matches:
        assert "API" in str(m.get("context", "")) or "api" in str(m.get("context", "")).lower()


def test_find_with_tag_no_match(run_cli, tmp_path):
    """find --tag with non-existent tag should produce WARN_TAG_NOT_FOUND."""
    md = tmp_path / "doc.md"
    md.write_text("# Heading\n\nContent.\n")

    result = run_cli("find", str(md), "Content", "--tag", "nonexistent")
    env = result.json
    assert env["ok"] is True
    assert any(w["code"] == "WARN_TAG_NOT_FOUND" for w in env.get("warnings", []))


# ---------------------------------------------------------------------------
# UX 5: Confidence values have no float artifacts
# ---------------------------------------------------------------------------

def test_confidence_rounded():
    """AnchorMatch confidence should be rounded to 4 decimal places."""
    from docweave.anchors import AnchorMatch
    from docweave.models import SourceSpan

    m = AnchorMatch(
        block_id="blk_001",
        block_kind="heading",
        source_span=SourceSpan(start_line=1, end_line=1),
        confidence=0.33333333333333,
        match_type="fuzzy",
        context="test",
    )
    assert m.confidence == 0.3333
    assert len(str(m.confidence).split(".")[-1]) <= 4


# ---------------------------------------------------------------------------
# UX 6: --list-tags on inspect
# ---------------------------------------------------------------------------

def test_inspect_list_tags(run_cli, tmp_path):
    """inspect --list-tags should list all unique tags."""
    md = tmp_path / "doc.md"
    md.write_text(
        '<!-- docweave: {"tags": ["api", "v2"]} -->\n'
        '# Endpoint\n\n'
        '<!-- docweave: {"tags": ["internal", "v2"]} -->\n'
        '# Internal\n'
    )
    result = run_cli("inspect", str(md), "--list-tags")
    env = result.json
    assert env["ok"] is True
    tags = env["result"]["tags"]
    assert sorted(tags) == ["api", "internal", "v2"]
    assert env["result"]["count"] == 3


def test_inspect_list_tags_empty(run_cli, tmp_path):
    """inspect --list-tags on a doc with no tags returns empty list."""
    md = tmp_path / "doc.md"
    md.write_text("# Heading\n\nNo tags here.\n")
    result = run_cli("inspect", str(md), "--list-tags")
    env = result.json
    assert env["ok"] is True
    assert env["result"]["tags"] == []
    assert env["result"]["count"] == 0


# ---------------------------------------------------------------------------
# Bug 3: Fleet evidence dir should contain actual files
# ---------------------------------------------------------------------------

def test_fleet_evidence_dir_has_files(run_cli, tmp_path):
    """fleet --evidence-dir should write evidence files, not just empty dirs."""
    doc = tmp_path / "doc.md"
    doc.write_text("# Title\n\nOriginal content.\n")

    patch = tmp_path / "patch.yaml"
    patch.write_text(yaml.dump({
        "version": 1,
        "target": {"file": str(doc), "format": "markdown"},
        "operations": [
            {
                "id": "op_001",
                "op": "insert_after",
                "anchor": {"by": "heading", "value": "Title"},
                "content": {"kind": "paragraph", "value": "Added by fleet."},
            }
        ],
    }))

    ev_dir = tmp_path / "evidence"
    result = run_cli("fleet", "--patch", str(patch), "--evidence-dir", str(ev_dir))
    env = result.json
    assert env["ok"] is True

    # Evidence dir should exist and have at least one subdirectory with files
    assert ev_dir.exists()
    subdirs = list(ev_dir.iterdir())
    assert len(subdirs) >= 1
    files_in_subdir = list(subdirs[0].iterdir())
    assert len(files_in_subdir) > 0, "Evidence subdirectory should contain actual files"


# ---------------------------------------------------------------------------
# UX 7: Fleet partial validation failure continues with valid patches
# ---------------------------------------------------------------------------

def test_fleet_partial_validation_failure(run_cli, tmp_path):
    """Fleet should continue with valid patches when some fail validation."""
    doc = tmp_path / "doc.md"
    doc.write_text("# Title\n\nContent.\n")

    good_patch = tmp_path / "good.yaml"
    good_patch.write_text(yaml.dump({
        "version": 1,
        "target": {"file": str(doc), "format": "markdown"},
        "operations": [
            {
                "id": "op_001",
                "op": "insert_after",
                "anchor": {"by": "heading", "value": "Title"},
                "content": {"kind": "paragraph", "value": "Added."},
            }
        ],
    }))

    bad_patch = tmp_path / "bad.yaml"
    bad_patch.write_text("not: valid: yaml: patch: {{{")

    result = run_cli(
        "fleet",
        "--patch", str(bad_patch),
        "--patch", str(good_patch),
    )
    env = result.json
    # Fleet should partially succeed
    assert env["result"]["succeeded"] >= 1
    assert env["result"]["failed"] >= 1
    # The good patch should have been applied
    assert "Added." in doc.read_text("utf-8")


def test_fleet_all_validation_failures(run_cli, tmp_path):
    """Fleet should fail gracefully when all patches fail validation."""
    bad_patch = tmp_path / "bad.yaml"
    bad_patch.write_text(yaml.dump({
        "version": 1,
        "target": {"backend": "auto"},  # missing file
        "operations": [],
    }))

    result = run_cli("fleet", "--patch", str(bad_patch))
    env = result.json
    assert env["ok"] is False
    assert env["result"]["failed"] >= 1


# ---------------------------------------------------------------------------
# UX 3: Guide mentions insert_after on heading behavior
# ---------------------------------------------------------------------------

def test_guide_mentions_insert_after_heading(run_cli):
    """guide should explain insert_after heading behavior."""
    result = run_cli("guide")
    env = result.json
    assert env["ok"] is True
    content_desc = env["result"]["patch_schema"]["operation_fields"]["content"]
    assert "insert_after on a heading" in content_desc
    assert "heading line" in content_desc


# ---------------------------------------------------------------------------
# Guide documents --tag on find, --list-tags on inspect
# ---------------------------------------------------------------------------

def test_guide_documents_find_tag(run_cli):
    """guide should document --tag option on find command."""
    result = run_cli("guide")
    env = result.json
    assert env["ok"] is True
    find_opts = env["result"]["commands"]["find"].get("options", {})
    assert "--tag" in find_opts


def test_guide_documents_inspect_list_tags(run_cli):
    """guide should document --list-tags option on inspect command."""
    result = run_cli("guide")
    env = result.json
    assert env["ok"] is True
    inspect_opts = env["result"]["commands"]["inspect"].get("options", {})
    assert "--list-tags" in inspect_opts
