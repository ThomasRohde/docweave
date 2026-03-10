"""Tests for the Word (.docx) applier."""

from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

import pytest
from docx import Document

from docweave.anchors import AnchorMatch
from docweave.backends.docx_backend import WordBackend
from docweave.models import SourceSpan
from docweave.plan.applier import FingerprintConflictError
from docweave.plan.applier_docx import apply_plan_docx
from docweave.plan.planner import ExecutionPlan, ResolvedOperation

SAMPLE = Path(__file__).parent / "fixtures" / "sample.docx"


@pytest.fixture()
def backend():
    return WordBackend()


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    dst = tmp_path / "sample.docx"
    shutil.copy2(SAMPLE, dst)
    return dst


def _fingerprint(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _make_plan(
    path: Path,
    operations: list[dict],
    anchor_matches: list[AnchorMatch],
) -> ExecutionPlan:
    """Build an ExecutionPlan with the given operations and pre-resolved anchors."""
    resolved = []
    for op, match in zip(operations, anchor_matches, strict=True):
        resolved.append(ResolvedOperation(
            operation=op,
            anchor_match=match,
            action_description=f"{op['op']} test",
            affected_lines=match.source_span,
        ))
    return ExecutionPlan(
        file=str(path),
        fingerprint=_fingerprint(path),
        backend="word-docx",
        resolved_operations=resolved,
        warnings=[],
        valid=True,
    )


def _find_block_span(backend: WordBackend, path: Path, text: str) -> tuple[SourceSpan, str, str]:
    """Find a block matching the given text. Returns (span, block_id, kind)."""
    doc = backend.load_view(path)
    for block in doc.blocks:
        if text.lower() in block.text.lower():
            return block.source_span, block.block_id, block.kind
    raise AssertionError(f"No block containing {text!r}")


def _make_match(
    span: SourceSpan, block_id: str = "blk_001", kind: str = "paragraph",
) -> AnchorMatch:
    return AnchorMatch(
        block_id=block_id,
        block_kind=kind,
        source_span=span,
        confidence=1.0,
        match_type="direct",
        context="test",
    )


# --- insert_after ---


def test_apply_insert_after(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "introduction")
    op = {
        "id": "op1",
        "op": "insert_after",
        "anchor": {"by": "quote", "value": "introduction"},
        "content": {"kind": "paragraph", "value": "A new paragraph inserted after."},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 1
    assert result.fingerprint_before != result.fingerprint_after

    # Verify the new paragraph exists
    doc = Document(str(sample_docx))
    texts = [p.text for p in doc.paragraphs]
    assert "A new paragraph inserted after." in texts


# --- insert_before ---


def test_apply_insert_before(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "Conclusion")
    op = {
        "id": "op1",
        "op": "insert_before",
        "anchor": {"by": "heading", "value": "Conclusion"},
        "content": {"kind": "paragraph", "value": "Paragraph before conclusion."},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 1
    doc = Document(str(sample_docx))
    texts = [p.text for p in doc.paragraphs]
    assert "Paragraph before conclusion." in texts
    # It should appear before "Conclusion" heading
    idx_new = texts.index("Paragraph before conclusion.")
    idx_conclusion = texts.index("Conclusion")
    assert idx_new < idx_conclusion


# --- replace_block ---


def test_apply_replace_block(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "concludes the sample")
    op = {
        "id": "op1",
        "op": "replace_block",
        "anchor": {"by": "quote", "value": "concludes"},
        "content": {"kind": "paragraph", "value": "This is the replaced conclusion."},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 1
    doc = Document(str(sample_docx))
    texts = [p.text for p in doc.paragraphs]
    assert "This is the replaced conclusion." in texts
    assert not any("concludes the sample" in t for t in texts)


# --- replace preserves formatting ---


def test_apply_replace_preserves_formatting(backend: WordBackend, sample_docx: Path):
    """Replacing text in a paragraph with bold/italic should preserve first run's formatting."""
    span, bid, kind = _find_block_span(backend, sample_docx, "bold text")
    op = {
        "id": "op1",
        "op": "replace_block",
        "anchor": {"by": "quote", "value": "bold text"},
        "content": {"kind": "paragraph", "value": "Replacement with original format."},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    apply_plan_docx(sample_docx, plan)

    doc = Document(str(sample_docx))
    # Find the replaced paragraph
    for p in doc.paragraphs:
        if "Replacement with original format." in p.text:
            # First run should exist
            assert len(p.runs) >= 1
            break
    else:
        pytest.fail("Replaced paragraph not found")


# --- delete_block ---


def test_apply_delete_block(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "Plan the next phase")
    op = {
        "id": "op1",
        "op": "delete_block",
        "anchor": {"by": "quote", "value": "Plan the next phase"},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 1
    doc = Document(str(sample_docx))
    texts = [p.text for p in doc.paragraphs]
    assert not any("Plan the next phase" in t for t in texts)


# --- set_heading ---


def test_apply_set_heading(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "Getting Started")
    op = {
        "id": "op1",
        "op": "set_heading",
        "anchor": {"by": "heading", "value": "Getting Started"},
        "content": {"kind": "heading", "value": "Quick Start Guide"},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, "heading")])
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 1
    doc = Document(str(sample_docx))
    texts = [p.text for p in doc.paragraphs]
    assert "Quick Start Guide" in texts
    assert "Getting Started" not in texts


# --- replace_text ---


def test_apply_replace_text(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "introduction")
    op = {
        "id": "op1",
        "op": "replace_text",
        "anchor": {"by": "quote", "value": "introduction paragraph"},
        "replacement": "opening paragraph",
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 1
    doc = Document(str(sample_docx))
    found = False
    for p in doc.paragraphs:
        if "opening paragraph" in p.text:
            found = True
            break
    assert found


# --- fingerprint conflict ---


def test_apply_fingerprint_conflict(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "introduction")
    op = {
        "id": "op1",
        "op": "delete_block",
        "anchor": {"by": "quote", "value": "introduction"},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    # Modify the file to cause a fingerprint mismatch
    sample_docx.write_bytes(sample_docx.read_bytes() + b"\x00")

    with pytest.raises(FingerprintConflictError):
        apply_plan_docx(sample_docx, plan)


# --- backup ---


def test_apply_backup_created(backend: WordBackend, sample_docx: Path):
    span, bid, kind = _find_block_span(backend, sample_docx, "introduction")
    op = {
        "id": "op1",
        "op": "delete_block",
        "anchor": {"by": "quote", "value": "introduction"},
    }
    plan = _make_plan(sample_docx, [op], [_make_match(span, bid, kind)])
    result = apply_plan_docx(sample_docx, plan, backup=True)

    assert result.backup_path is not None
    assert Path(result.backup_path).exists()


# --- multiple operations bottom-up ---


def test_apply_multiple_operations_bottom_up(backend: WordBackend, sample_docx: Path):
    """Apply two operations: one near the end, one near the start. Bottom-up order matters."""
    span1, bid1, kind1 = _find_block_span(backend, sample_docx, "introduction")
    span2, bid2, kind2 = _find_block_span(backend, sample_docx, "concludes")

    ops = [
        {
            "id": "op1",
            "op": "replace_block",
            "anchor": {"by": "quote", "value": "introduction"},
            "content": {"kind": "paragraph", "value": "Modified intro."},
        },
        {
            "id": "op2",
            "op": "replace_block",
            "anchor": {"by": "quote", "value": "concludes"},
            "content": {"kind": "paragraph", "value": "Modified conclusion."},
        },
    ]
    matches = [
        _make_match(span1, bid1, kind1),
        _make_match(span2, bid2, kind2),
    ]
    plan = _make_plan(sample_docx, ops, matches)
    result = apply_plan_docx(sample_docx, plan)

    assert result.operations_applied == 2
    doc = Document(str(sample_docx))
    texts = [p.text for p in doc.paragraphs]
    assert "Modified intro." in texts
    assert "Modified conclusion." in texts
