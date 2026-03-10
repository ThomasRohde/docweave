"""Helper script to generate tests/fixtures/sample.docx programmatically."""

from pathlib import Path

from docx import Document


def create_sample_docx(output_path: Path) -> None:
    """Create a sample .docx with headings, paragraphs, lists, table, and formatting."""
    doc = Document()

    # Title (Heading 1)
    doc.add_heading("Sample Document", level=1)

    # Introduction paragraph
    doc.add_paragraph("This is the introduction paragraph with some plain text.")

    # Section 1 (Heading 2)
    doc.add_heading("Getting Started", level=2)

    # Paragraph with mixed formatting
    p = doc.add_paragraph()
    p.add_run("This paragraph has ")
    bold_run = p.add_run("bold text")
    bold_run.bold = True
    p.add_run(" and ")
    italic_run = p.add_run("italic text")
    italic_run.italic = True
    p.add_run(" mixed in.")

    # Bulleted list
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    doc.add_paragraph("Third item", style="List Bullet")

    # Section 2 (Heading 2)
    doc.add_heading("Data Overview", level=2)

    doc.add_paragraph("The following table shows key metrics:")

    # Table (3 columns x 3 rows including header)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    # Header row
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Status"
    # Data rows
    table.rows[1].cells[0].text = "Performance"
    table.rows[1].cells[1].text = "95%"
    table.rows[1].cells[2].text = "Good"
    table.rows[2].cells[0].text = "Uptime"
    table.rows[2].cells[1].text = "99.9%"
    table.rows[2].cells[2].text = "Excellent"

    # Section 3 (Heading 2)
    doc.add_heading("Conclusion", level=2)

    doc.add_paragraph("This concludes the sample document.")

    # Subsection (Heading 3)
    doc.add_heading("Next Steps", level=3)

    doc.add_paragraph("Plan the next phase of the project.")

    doc.save(str(output_path))


if __name__ == "__main__":
    out = Path(__file__).parent / "fixtures" / "sample.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    create_sample_docx(out)
    print(f"Created {out}")
